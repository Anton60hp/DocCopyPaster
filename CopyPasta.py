import os
import docx
from docx.shared import Pt
import pandas as pd
import configparser


config = configparser.ConfigParser()
config.read_file(open("config.ini",encoding='utf8'))

# Word config import
DOC_FILE_NAME = config['Word']['DOC_FILE_NAME']
CODE_PHRASE = config['Word']['CODE_PHRASE']

# Excel config import
EXCEL_FILE_NAME = config['Excel']['EXCEL_FILE_NAME']
COLOMN = int(config['Excel']['COLOMN']) - 1
SKIP_LINES = int(config['Excel']['START_LINE'])

# Result config import
DIRECTORY_NAME = config['Result']['DIRECTORY_NAME']
FONT_SIZE = int(config['Result']['FONT_SIZE'])
FILE_NAME = config['Result']['FILE_NAME']

# Dataframe creation
df1 = pd.read_excel(EXCEL_FILE_NAME, header=None, index_col=None, skiprows=lambda x: x in range(SKIP_LINES))
df1 = df1[COLOMN]

new_doc = docx.Document(DOC_FILE_NAME)
paragraphs_to_change = []
for i in range(len(new_doc.paragraphs)):
    if CODE_PHRASE in new_doc.paragraphs[i].text:
        paragraphs_to_change.append(i)

if not paragraphs_to_change:
    print('Кодовая фраза в тексте не найдена.')
else:
    print('Кодовая фраза найдена, начинаю создание файлов')

    # Directory creation
    if not os.path.exists(DIRECTORY_NAME):
        os.mkdir(DIRECTORY_NAME)
    os.chdir(DIRECTORY_NAME)

    for i in range(len(df1)):
        os.chdir("../")
        new_doc = docx.Document(DOC_FILE_NAME)
        for k in paragraphs_to_change:
            new_doc.paragraphs[k].text = new_doc.paragraphs[k].text.replace(CODE_PHRASE, df1.iloc[i])
            new_doc.paragraphs[k].style.font.size = Pt(FONT_SIZE)
        os.chdir("result")
        new_doc.save(f'{FILE_NAME}{i+1}.docx')

    print(f'Done. {len(df1)} files created.')


    

